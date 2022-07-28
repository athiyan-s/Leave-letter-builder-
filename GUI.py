from os import register_at_fork
from tkinter import *
from tkinter import ttk
import replace
import re
from docx import Document
document = Document()
doc = Document('/home/athiyan/python codes/LETTER/source/Leave Letter.docx')

root = Tk()

def printValue():
    pname = name_field.get() 
    pcourse = course_field.get()    
    pyear = year_field.get()
    pm_s=m_s_no_field.get()
    preason=reason_field.get()
    pdate=date_field.get()
    

    chg_find=()
    chg_replace=()





    
    find_value=['!!!','@@@','###','&&&','$$$',"~~~"]
    replace_value=[pname,pcourse,pyear,pm_s,preason,pdate]
    i=0
    a=0
    while i<6:
        
        regex1 = re.compile(find_value[i])
        chg_find=regex1
        print(chg_find)
        i=i+1
        replace1 = (replace_value[a])
        chg_replace=replace1
        print(chg_replace)
        a=a+1
        print("done 1")
        filename = "Leave Letter.docx"
        doc = Document(filename)
        replace.docx_replace_regex(doc, chg_find , chg_replace)
        print("REPLACED")
        doc.save("/home/athiyan/python codes/LETTER/hello_world.docx")
        print("done")
    #doc.save("/home/athiyan/python codes/LETTER/hello_world.docx")
    print("done and completed")

#find_value=['!!!','@@@','###','&&&','$$$',"~~~"]
#repalce_value=['name','dept','year','m_s','reason','date']

root.configure(background='light green')
root.title("registration form")
root.geometry("500x300")
heading = Label(root, text="Form", bg="light green")
name = Label(root, text="Name", bg="light green").grid(row = 1,column = 0)
course = Label(root, text="Course", bg="light green").grid(row = 2,column = 0)
year = Label(root, text="year", bg="light green").grid(row = 3,column = 0)
m_s = Label(root, text="SIR OR MAM", bg="light green").grid(row = 4,column = 0)
reason = Label(root, text="REASON", bg="light green").grid(row = 5,column = 0)
date = Label(root, text="DATE ON", bg="light green").grid(row = 6,column = 0)

name_field = Entry(root)
course_field = Entry(root)
year_field = Entry(root)
m_s_no_field = Entry(root)
reason_field = Entry(root)
date_field = Entry(root)


name_field.grid(row=1, column=1, ipadx="100")
course_field.grid(row=2, column=1, ipadx="100")
year_field.grid(row=3, column=1, ipadx="100")
m_s_no_field.grid(row=4, column=1, ipadx="100")
reason_field.grid(row=5, column=1, ipadx="100")
date_field.grid(row=6, column=1, ipadx="100")

submit = Button(root, text="Submit",command=printValue, fg="Black",bg="Red")
submit.grid(row=8, column=1)


doc.save("/home/athiyan/python codes/LETTER/hello_world.docx")
root.mainloop()
