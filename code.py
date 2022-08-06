import replace
import re
from docx import Document

document = Document()
doc = Document('Leave Letter.docx')
    
pname =input("name: ")
pemail=input("Your Email: ")  
pcourse =input("course: ")    
pyear = input("year: ")   
pm_s=input("to ") 
pmam_sir=input("mam or sir: ")
preason=input("reason: ")   
pdate=input("date: ")   

chg_find=()
chg_replace=()

find_value=['!!!','@@@','###','&&&','^^^','<<<',"~~~"]
replace_value=[pname,pcourse,pyear,pm_s,pmam_sir,preason,pdate]
i=0
a=0
while i<7:
    
    regex1 = re.compile(find_value[i])
    chg_find=regex1
    print(chg_find)
    i=i+1
    replace1 = (replace_value[a])
    chg_replace=replace1
    print(chg_replace)
    a=a+1
    print("done 1")
    filename = "Leave Letter.docx"
    doc = Document(filename)
    replace.docx_replace_regex(doc, chg_find , chg_replace)
    print("REPLACED")
    doc.save("Leave Letter.docx")
    print("done")
#doc.save("/home/athiyan/python codes/LETTER/hello_world.docx")
print("done and completed")
